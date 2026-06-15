import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.81)
        section.right_margin = Cm(2.54)

def setup_styles(doc):
    # Default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5
    
    # Adjust headings to look like bold Normal text but size 14
    for i in range(1, 4):
        try:
            hx_style = doc.styles[f'Heading {i}']
            hx_font = hx_style.font
            hx_font.name = 'Times New Roman'
            hx_font.size = Pt(14)
            hx_font.bold = True
            hx_font.color.rgb = None # default black
            hx_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hx_style.paragraph_format.line_spacing = 1.5
        except KeyError:
            pass

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = None

def generate_technical_doc():
    doc = Document()
    set_margins(doc)
    setup_styles(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documentación del Proyecto AEROPUERTO_V2")
    run.font.size = Pt(14)
    run.font.bold = True
    
    add_heading(doc, "Índice", level=1)
    doc.add_paragraph("1. Objetivos\n2. Alcances y Limitaciones\n3. Resumen\n4. Introducción\n5. Marco Teórico\n6. Lenguaje, Base de Datos y Herramientas\n7. Funcionamiento del Proyecto\n8. Desarrollo del Proyecto\n9. Resultados\n10. Conclusión y Recomendaciones\n11. Fuentes de Información\n12. Anexos")
    doc.add_page_break()
    
    add_heading(doc, "1. Objetivos", level=1)
    doc.add_paragraph("Objetivo General:\nDesarrollar e implementar un sistema integral de reservas de vuelos (AEROPUERTO_V2) que permita a los pasajeros buscar, reservar y pagar vuelos de manera eficiente, así como brindar a los administradores herramientas para la gestión de vuelos, usuarios y reservas.\n\nObjetivos Específicos:\n- Implementar un sistema de autenticación y autorización seguro para diferentes roles (pasajeros, agentes, administradores).\n- Desarrollar una interfaz gráfica moderna, amigable y responsiva.\n- Integrar una pasarela de pagos segura (Stripe) para el procesamiento de transacciones financieras.\n- Diseñar y administrar una base de datos relacional para el manejo eficiente y estructurado de la información.")

    add_heading(doc, "2. Alcances y Limitaciones", level=1)
    doc.add_paragraph("Alcances:\nEl proyecto cubre la gestión completa del ciclo de vida de una reserva de vuelo, desde la búsqueda de disponibilidad, selección de asientos, hasta el pago y confirmación de la reserva. Incluye notificaciones y perfiles de usuario.\n\nLimitaciones:\nEl sistema se limita actualmente a operaciones simuladas de pasarela de pagos en entornos de prueba y no contempla integraciones en tiempo real con sistemas de radar o control de tráfico aéreo físico.")

    add_heading(doc, "3. Resumen", level=1)
    doc.add_paragraph("El presente trabajo detalla el proceso de desarrollo e implementación del proyecto AEROPUERTO_V2. El sistema es una aplicación web full-stack orientada a la gestión de reservas de vuelos. A través del documento se exploran los aspectos teóricos, tecnológicos y de funcionamiento, evidenciando un diseño de arquitectura robusto y una implementación basada en prácticas modernas de desarrollo de software.")

    add_heading(doc, "4. Introducción", level=1)
    doc.add_paragraph("La modernización de los sistemas aeroportuarios es fundamental para brindar un mejor servicio. AEROPUERTO_V2 nace como una solución integral que digitaliza y automatiza las reservas de vuelos. Este proyecto se ha desarrollado considerando aspectos críticos como la seguridad, usabilidad y escalabilidad. A continuación, se presenta la documentación detallada que abarca desde la conceptualización hasta los resultados obtenidos.")

    add_heading(doc, "5. Marco Teórico", level=1)
    doc.add_paragraph("El desarrollo de aplicaciones web modernas se basa en la separación de responsabilidades a través de arquitecturas cliente-servidor. Se implementan principios de diseño REST para la comunicación y sistemas ORM (Object-Relational Mapping) para la interacción con la base de datos, lo cual optimiza la gestión y reduce la complejidad del código SQL directo. Además, se aplican protocolos criptográficos para el almacenamiento seguro de contraseñas (Bcrypt) y tokens de acceso (JWT).")

    add_heading(doc, "6. Lenguaje, Base de Datos y Herramientas", level=1)
    doc.add_paragraph("El ecosistema tecnológico utilizado incluye:\n- Frontend: Desarrollado en TypeScript utilizando el framework React/Next.js o Vite.\n- Backend: API RESTful construida con Node.js, Express y TypeScript.\n- Base de Datos: PostgreSQL gestionado mediante el ORM Prisma.\n- Pagos: Integración con Stripe.\n- Infraestructura y Despliegue: Docker para contenedores y Render para despliegue de servicios.\n- Pruebas: Jest y Supertest.")

    add_heading(doc, "7. Funcionamiento del Proyecto", level=1)
    doc.add_paragraph("El sistema permite a los usuarios no registrados visualizar vuelos disponibles filtrando por origen y destino. Una vez registrados, pueden reservar asientos específicos y procesar su pago. El backend procesa estas solicitudes, calcula precios aplicando ofertas aplicables, y registra la transacción en PostgreSQL. Las notificaciones se generan automáticamente para informar al usuario sobre el estado de su reserva y pago.")

    add_heading(doc, "8. Desarrollo del Proyecto", level=1)
    doc.add_paragraph("El desarrollo siguió una metodología iterativa e incremental. Se comenzó por el diseño de la base de datos y la creación del esquema de Prisma. Posteriormente, se implementaron los controladores y rutas en el backend utilizando Node.js y Express, garantizando la validación de datos con Zod. En paralelo, se construyeron los componentes de la interfaz de usuario en React, y finalmente se realizó la integración de ambos, seguida por pruebas y ajustes de seguridad.")

    add_heading(doc, "9. Resultados", level=1)
    doc.add_paragraph("Como resultado, se obtuvo una plataforma web completamente funcional y desplegada. El sistema de reservas procesa exitosamente los flujos de usuarios, el control de acceso basado en roles funciona correctamente, y la integración de pagos procesa de manera simulada las transacciones financieras con alta fiabilidad.")

    add_heading(doc, "10. Conclusión y Recomendaciones", level=1)
    doc.add_paragraph("Conclusión:\nAEROPUERTO_V2 ha cumplido satisfactoriamente con los requerimientos planteados, demostrando la eficacia de combinar tecnologías modernas como Node.js, React y PostgreSQL para crear soluciones robustas y escalables.\n\nRecomendaciones:\nSe recomienda en futuras versiones integrar autenticación de dos factores (2FA), así como expandir la cobertura de pruebas de integración a más flujos críticos para asegurar un mantenimiento más ágil.")

    add_heading(doc, "11. Fuentes de Información", level=1)
    doc.add_paragraph("- Documentación oficial de Node.js.\n- Documentación oficial de React.\n- Manual de PostgreSQL.\n- Guías de Prisma ORM.\n- Referencias de la API de Stripe.")

    add_heading(doc, "12. Anexos", level=1)
    doc.add_paragraph("A continuación, se presentan los diagramas que ilustran la estructura y el flujo de la aplicación.")
    
    if os.path.exists("er_diagram.png"):
        doc.add_paragraph("Anexo 1: Diagrama Entidad-Relación")
        doc.add_picture("er_diagram.png", width=Inches(6))
    
    if os.path.exists("sequence_diagram.png"):
        doc.add_paragraph("Anexo 2: Diagrama de Secuencia")
        doc.add_picture("sequence_diagram.png", width=Inches(6))

    doc.save("Documentacion_Tecnica.docx")

def generate_user_manual():
    doc = Document()
    set_margins(doc)
    setup_styles(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual de Usuario - AEROPUERTO_V2")
    run.font.size = Pt(14)
    run.font.bold = True
    
    add_heading(doc, "Índice", level=1)
    doc.add_paragraph("1. Introducción\n2. Requisitos del Sistema\n3. Registro e Inicio de Sesión\n4. Búsqueda y Reserva de Vuelos\n5. Pagos\n6. Gestión de Perfil y Reservas\n7. Soporte")
    doc.add_page_break()
    
    add_heading(doc, "1. Introducción", level=1)
    doc.add_paragraph("El presente Manual de Usuario tiene como propósito guiar a los usuarios del sistema AEROPUERTO_V2 a través de las diferentes funcionalidades de la plataforma. Proporciona instrucciones paso a paso para asegurar una experiencia fluida desde la creación de una cuenta hasta la gestión completa de sus reservas de vuelo.")

    add_heading(doc, "2. Requisitos del Sistema", level=1)
    doc.add_paragraph("Para utilizar la plataforma AEROPUERTO_V2, únicamente se necesita:\n- Un dispositivo con conexión a internet estable.\n- Un navegador web moderno y actualizado (Google Chrome, Mozilla Firefox, Safari, Edge).\n- Una cuenta de correo electrónico válida para el registro y recepción de notificaciones.")

    add_heading(doc, "3. Registro e Inicio de Sesión", level=1)
    doc.add_paragraph("Registro:\n1. Acceda a la página principal del sistema.\n2. Haga clic en el botón 'Registrarse'.\n3. Complete el formulario con su nombre completo, correo electrónico y contraseña.\n4. Haga clic en 'Crear Cuenta'.\n\nInicio de Sesión:\n1. Desde la página principal, haga clic en 'Iniciar Sesión'.\n2. Ingrese su correo electrónico y contraseña.\n3. Haga clic en 'Acceder'.")

    add_heading(doc, "4. Búsqueda y Reserva de Vuelos", level=1)
    doc.add_paragraph("Búsqueda:\n1. En la página de inicio, utilice el buscador ingresando la ciudad de origen y la ciudad de destino.\n2. Seleccione la fecha de su viaje.\n3. Haga clic en 'Buscar Vuelos'. El sistema listará las opciones disponibles.\n\nReserva:\n1. En la lista de resultados, seleccione el vuelo de su preferencia haciendo clic en 'Ver Detalles'.\n2. Escoja el asiento disponible que desee.\n3. Revise el desglose del precio y haga clic en 'Reservar'.")

    add_heading(doc, "5. Pagos", level=1)
    doc.add_paragraph("Una vez confirmada su reserva, el sistema lo redirigirá a la pasarela de pagos:\n1. Revise el monto total a pagar.\n2. Ingrese los detalles de su método de pago (tarjeta de crédito/débito).\n3. Haga clic en 'Pagar'.\n4. El sistema le mostrará una pantalla de confirmación (Payment Success) y enviará un recibo a su correo electrónico.")

    add_heading(doc, "6. Gestión de Perfil y Reservas", level=1)
    doc.add_paragraph("Para administrar su cuenta:\n1. Inicie sesión y diríjase a 'Mi Perfil'.\n2. En esta sección podrá actualizar sus datos personales.\n3. En la sección 'Mis Reservas', podrá ver el historial de sus vuelos, el estado de los mismos y descargar comprobantes de pago.")

    add_heading(doc, "7. Soporte", level=1)
    doc.add_paragraph("En caso de presentarse algún error o inconveniente durante el uso de la plataforma, por favor contacte a nuestro equipo de soporte enviando un correo electrónico detallando el problema. Estaremos encantados de asistirle para asegurar que su experiencia sea óptima.")

    doc.save("Manual_Usuario.docx")

if __name__ == '__main__':
    generate_technical_doc()
    generate_user_manual()
    print("Documents generated successfully.")
